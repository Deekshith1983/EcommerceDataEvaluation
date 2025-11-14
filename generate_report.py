import sqlite3
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import matplotlib.pyplot as plt
import seaborn as sns

# Connect to the database
conn = sqlite3.connect('ecom.db')

# SQL query to join all tables
query = """
SELECT 
    c.name AS customer_name,
    c.city,
    o.order_id,
    o.order_date,
    p.product_name,
    oi.quantity,
    oi.item_price,
    o.total_amount,
    s.shipping_status,
    s.delivery_date
FROM customers c
INNER JOIN orders o ON c.customer_id = o.customer_id
INNER JOIN order_items oi ON o.order_id = oi.order_id
INNER JOIN products p ON oi.product_id = p.product_id
INNER JOIN shipping s ON o.order_id = s.order_id
ORDER BY o.order_id, oi.order_item_id
"""

print("Executing SQL query to generate comprehensive e-commerce report...\n")
print("="*80)
print("SQL QUERY:")
print("="*80)
print(query)
print("="*80 + "\n")

# Execute query and fetch results
df = pd.read_sql_query(query, conn)

print(f"Query executed successfully!")
print(f"Total records retrieved: {len(df)}\n")

# Display first 20 rows
print("="*80)
print("SAMPLE REPORT (First 20 rows):")
print("="*80)
print(df.head(20).to_string(index=False))

print("\n" + "="*80)
print("REPORT STATISTICS:")
print("="*80)
print(f"Total Records: {len(df)}")
print(f"Unique Customers: {df['customer_name'].nunique()}")
print(f"Unique Orders: {df['order_id'].nunique()}")
print(f"Unique Products: {df['product_name'].nunique()}")
print(f"\nShipping Status Distribution:")
print(df['shipping_status'].value_counts().to_string())

# Save to CSV
output_file = 'ecommerce_report.csv'
df.to_csv(output_file, index=False)
print(f"\n✓ Full report saved to CSV: {output_file}")

# Generate Visualizations
print("\n" + "="*80)
print("GENERATING VISUALIZATIONS...")
print("="*80)

# Set style
sns.set_style("whitegrid")
plt.rcParams['figure.figsize'] = (10, 6)

# Visualization 1: Shipping Status Distribution (Pie Chart)
plt.figure(figsize=(8, 8))
status_counts = df['shipping_status'].value_counts()
colors_pie = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#FFA07A', '#98D8C8']
plt.pie(status_counts.values, labels=status_counts.index, autopct='%1.1f%%', 
        startangle=90, colors=colors_pie, explode=[0.05]*len(status_counts))
plt.title('Shipping Status Distribution', fontsize=16, fontweight='bold', pad=20)
plt.axis('equal')
plt.tight_layout()
plt.savefig('shipping_status_chart.png', dpi=300, bbox_inches='tight')
plt.close()
print("✓ Generated: shipping_status_chart.png")

# Visualization 2: Top 10 Products by Revenue (Bar Chart)
plt.figure(figsize=(12, 8))
product_revenue = df.groupby('product_name')['item_price'].sum().sort_values(ascending=False).head(10)
bars = plt.barh(range(len(product_revenue)), product_revenue.values, color='#1f4788')
plt.yticks(range(len(product_revenue)), product_revenue.index)
plt.xlabel('Total Revenue (₹)', fontsize=12, fontweight='bold')
plt.ylabel('Product Name', fontsize=12, fontweight='bold')
plt.title('Top 10 Products by Revenue', fontsize=16, fontweight='bold', pad=20)
plt.gca().invert_yaxis()

# Add value labels on bars
for i, (idx, value) in enumerate(product_revenue.items()):
    plt.text(value, i, f' ₹{value:,.0f}', va='center', fontsize=10)

plt.tight_layout()
plt.savefig('top_products_chart.png', dpi=300, bbox_inches='tight')
plt.close()
print("✓ Generated: top_products_chart.png")

print("\n" + "="*80)
print("GENERATING DOCX REPORT...")
print("="*80)

doc = Document()

# Add title
title = doc.add_heading('E-Commerce Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add report metadata
doc.add_heading('Report Statistics', level=1)
stats_para = doc.add_paragraph()
stats_para.add_run(f'Total Records: {len(df)}\n').bold = True
stats_para.add_run(f'Unique Customers: {df["customer_name"].nunique()}\n')
stats_para.add_run(f'Unique Orders: {df["order_id"].nunique()}\n')
stats_para.add_run(f'Unique Products: {df["product_name"].nunique()}\n')

# Add shipping status distribution
doc.add_heading('Shipping Status Distribution', level=1)
status_counts = df['shipping_status'].value_counts()
for status, count in status_counts.items():
    doc.add_paragraph(f'{status}: {count}', style='List Bullet')

# Add visualizations to DOCX
doc.add_page_break()
doc.add_heading('Data Visualizations', level=1)

# Add shipping status chart
doc.add_heading('1. Shipping Status Distribution', level=2)
doc.add_picture('shipping_status_chart.png', width=Inches(6))

doc.add_paragraph()  # Add spacing

# Add top products chart
doc.add_heading('2. Top 10 Products by Revenue', level=2)
doc.add_picture('top_products_chart.png', width=Inches(6))

doc.add_page_break()

# Add sample data table (first 50 rows)
doc.add_heading('Sample Data (First 50 Records)', level=1)

# Create table
table = doc.add_table(rows=1, cols=10)
table.style = 'Light Grid Accent 1'

# Add headers
header_cells = table.rows[0].cells
headers = ['Customer Name', 'City', 'Order ID', 'Order Date', 'Product Name', 
           'Quantity', 'Item Price', 'Total Amount', 'Shipping Status', 'Delivery Date']
for i, header in enumerate(headers):
    header_cells[i].text = header
    header_cells[i].paragraphs[0].runs[0].font.bold = True

# Add data rows (first 50)
for idx, row in df.head(50).iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(row['customer_name'])
    row_cells[1].text = str(row['city'])
    row_cells[2].text = str(row['order_id'])
    row_cells[3].text = str(row['order_date'])
    row_cells[4].text = str(row['product_name'])
    row_cells[5].text = str(row['quantity'])
    row_cells[6].text = f"₹{row['item_price']:.2f}"
    row_cells[7].text = f"₹{row['total_amount']:.2f}"
    row_cells[8].text = str(row['shipping_status'])
    row_cells[9].text = str(row['delivery_date'])

docx_file = 'ecommerce_report.docx'
doc.save(docx_file)
print(f"✓ Full report saved to DOCX: {docx_file}")

# Generate PDF Report
print("\n" + "="*80)
print("GENERATING PDF REPORT...")
print("="*80)

pdf_file = 'ecommerce_report.pdf'
pdf = SimpleDocTemplate(pdf_file, pagesize=A4)
elements = []

# Styles
styles = getSampleStyleSheet()
title_style = ParagraphStyle(
    'CustomTitle',
    parent=styles['Heading1'],
    fontSize=24,
    textColor=colors.HexColor('#1f4788'),
    spaceAfter=30,
    alignment=1  # Center
)

heading_style = ParagraphStyle(
    'CustomHeading',
    parent=styles['Heading2'],
    fontSize=14,
    textColor=colors.HexColor('#1f4788'),
    spaceAfter=12,
)

# Add title
elements.append(Paragraph("E-Commerce Report", title_style))
elements.append(Spacer(1, 0.3*inch))

# Add statistics
elements.append(Paragraph("Report Statistics", heading_style))
stats_text = f"""
<b>Total Records:</b> {len(df)}<br/>
<b>Unique Customers:</b> {df['customer_name'].nunique()}<br/>
<b>Unique Orders:</b> {df['order_id'].nunique()}<br/>
<b>Unique Products:</b> {df['product_name'].nunique()}
"""
elements.append(Paragraph(stats_text, styles['Normal']))
elements.append(Spacer(1, 0.2*inch))

# Add shipping status
elements.append(Paragraph("Shipping Status Distribution", heading_style))
status_text = "<br/>".join([f"<b>{status}:</b> {count}" for status, count in status_counts.items()])
elements.append(Paragraph(status_text, styles['Normal']))
elements.append(Spacer(1, 0.3*inch))

# Add visualizations to PDF
elements.append(PageBreak())
elements.append(Paragraph("Data Visualizations", title_style))
elements.append(Spacer(1, 0.2*inch))

# Add shipping status chart
elements.append(Paragraph("1. Shipping Status Distribution", heading_style))
elements.append(Spacer(1, 0.1*inch))
elements.append(Image('shipping_status_chart.png', width=5*inch, height=5*inch))
elements.append(Spacer(1, 0.3*inch))

# Add top products chart
elements.append(Paragraph("2. Top 10 Products by Revenue", heading_style))
elements.append(Spacer(1, 0.1*inch))
elements.append(Image('top_products_chart.png', width=6*inch, height=4*inch))
elements.append(Spacer(1, 0.3*inch))

elements.append(PageBreak())

# Add data table (first 30 rows for PDF)
elements.append(Paragraph("Sample Data (First 30 Records)", heading_style))
elements.append(Spacer(1, 0.1*inch))

# Prepare table data
table_data = [['Customer', 'City', 'Order ID', 'Date', 'Product', 
               'Qty', 'Price', 'Total', 'Status', 'Delivery']]

for idx, row in df.head(30).iterrows():
    table_data.append([
        str(row['customer_name'])[:15],  # Truncate long names
        str(row['city'])[:10],
        str(row['order_id']),
        str(row['order_date']),
        str(row['product_name'])[:15],
        str(row['quantity']),
        f"₹{row['item_price']:.0f}",
        f"₹{row['total_amount']:.0f}",
        str(row['shipping_status'])[:10],
        str(row['delivery_date'])
    ])

# Create table
pdf_table = Table(table_data, repeatRows=1)
pdf_table.setStyle(TableStyle([
    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4788')),
    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
    ('FONTSIZE', (0, 0), (-1, 0), 8),
    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
    ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ('FONTSIZE', (0, 1), (-1, -1), 6),
    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
]))

elements.append(pdf_table)

# Build PDF
pdf.build(elements)
print(f"✓ Full report saved to PDF: {pdf_file}")

print("\n" + "="*80)
print("REPORT GENERATION COMPLETE!")
print("="*80)
print(f"📄 CSV Report: {output_file}")
print(f"📄 DOCX Report: {docx_file}")
print(f"📄 PDF Report: {pdf_file}")
print(f"📊 Visualizations:")
print(f"   - shipping_status_chart.png")
print(f"   - top_products_chart.png")
print("="*80)



# Close connection
conn.close()
print("\nDatabase connection closed.")
